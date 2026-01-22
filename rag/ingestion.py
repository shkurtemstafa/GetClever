"""Document ingestion pipeline - loads, cleans, chunks and processes documents."""

import os
import hashlib
from typing import List, Dict, Any
from pathlib import Path

import PyPDF2
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup

from langchain_core.documents import Document as LangchainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter


# Configuration
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
DOCUMENTS_DIR = "./dataset"


class DocumentProcessor:
    """Handles document loading, cleaning, and chunking."""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def load_documents(self, directory: str = DOCUMENTS_DIR) -> List[LangchainDocument]:
        """Load all documents from the specified directory."""
        documents = []
        
        for file_path in Path(directory).rglob("*"):
            if file_path.is_file() and file_path.name != "GetClever.png":
                try:
                    doc_content = self._load_single_document(file_path)
                    if doc_content:
                        documents.extend(doc_content)
                except Exception as e:
                    print(f"Error loading {file_path}: {e}")
        
        return documents
    
    def _load_single_document(self, file_path: Path) -> List[LangchainDocument]:
        """Load a single document based on its file type."""
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            return self._load_pdf(file_path)
        elif file_extension == '.docx':
            return self._load_docx(file_path)
        elif file_extension == '.md':
            return self._load_markdown(file_path)
        elif file_extension == '.txt':
            return self._load_text(file_path)
        else:
            print(f"Unsupported file type: {file_extension}")
            return []
    
    def _load_pdf(self, file_path: Path) -> List[LangchainDocument]:
        """Load PDF document."""
        documents = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    cleaned_text = self._clean_text(text)
                    chunks = self.text_splitter.split_text(cleaned_text)
                    
                    for chunk_idx, chunk in enumerate(chunks):
                        doc = LangchainDocument(
                            page_content=chunk,
                            metadata={
                                "source": str(file_path.name),
                                "page": page_num + 1,
                                "chunk": chunk_idx + 1,
                                "file_type": "pdf",
                                "doc_id": self._generate_doc_id(file_path, page_num, chunk_idx)
                            }
                        )
                        documents.append(doc)
        
        return documents
    
    def _load_docx(self, file_path: Path) -> List[LangchainDocument]:
        """Load DOCX document."""
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        cleaned_text = self._clean_text(text)
        chunks = self.text_splitter.split_text(cleaned_text)
        
        documents = []
        for chunk_idx, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": str(file_path.name),
                    "chunk": chunk_idx + 1,
                    "file_type": "docx",
                    "doc_id": self._generate_doc_id(file_path, 0, chunk_idx)
                }
            )
            documents.append(doc)
        
        return documents
    
    def _load_markdown(self, file_path: Path) -> List[LangchainDocument]:
        """Load Markdown document."""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
        
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()
        
        cleaned_text = self._clean_text(text)
        chunks = self.text_splitter.split_text(cleaned_text)
        
        documents = []
        for chunk_idx, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": str(file_path.name),
                    "chunk": chunk_idx + 1,
                    "file_type": "markdown",
                    "doc_id": self._generate_doc_id(file_path, 0, chunk_idx)
                }
            )
            documents.append(doc)
        
        return documents
    
    def _load_text(self, file_path: Path) -> List[LangchainDocument]:
        """Load plain text document."""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        cleaned_text = self._clean_text(text)
        chunks = self.text_splitter.split_text(cleaned_text)
        
        documents = []
        for chunk_idx, chunk in enumerate(chunks):
            doc = LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": str(file_path.name),
                    "chunk": chunk_idx + 1,
                    "file_type": "txt",
                    "doc_id": self._generate_doc_id(file_path, 0, chunk_idx)
                }
            )
            documents.append(doc)
        
        return documents
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        text = " ".join(text.split())
        text = text.replace('\x00', '')
        return text.strip()
    
    def _generate_doc_id(self, file_path: Path, page: int, chunk: int) -> str:
        """Generate unique document ID."""
        content = f"{file_path.name}_{page}_{chunk}"
        return hashlib.md5(content.encode()).hexdigest()[:12]
    
    def get_document_stats(self, documents: List[LangchainDocument]) -> Dict[str, Any]:
        """Get statistics about processed documents."""
        if not documents:
            return {}
        
        stats = {
            "total_documents": len(documents),
            "total_characters": sum(len(doc.page_content) for doc in documents),
            "avg_chunk_size": sum(len(doc.page_content) for doc in documents) / len(documents),
            "sources": list(set(doc.metadata["source"] for doc in documents)),
            "file_types": list(set(doc.metadata["file_type"] for doc in documents))
        }
        
        return stats